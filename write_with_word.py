from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sqlite3

def fetch_data():
    conn = sqlite3.connect('material_company.db')
    cur = conn.cursor()
    cur.execute('SELECT * FROM material')
    material = cur.fetchall()
    cur.execute('SELECT * FROM company')
    company = cur.fetchall()
    cur.execute('SELECT * FROM blueprint')
    blueprint = cur.fetchall()
    cur.execute('SELECT * FROM orderlist')
    orderlist = cur.fetchall()
    conn.close()
    return material, company, blueprint, orderlist

def search_company_info(company_name):
    conn = sqlite3.connect('material_company.db')
    cur = conn.cursor()
    cur.execute('''
                SELECT *
                FROM company
                WHERE cName = ?;''', (company_name,))
    company_info = cur.fetchall()
    conn.close()
    return company_info

def search_blueprint_info(blueprint_id):
    conn = sqlite3.connect('material_company.db')
    cur = conn.cursor()
    cur.execute('''
                SELECT *
                FROM blueprint
                WHERE bid = ?;''', (blueprint_id,))
    blueprint_info = cur.fetchall()
    conn.close()
    return blueprint_info

material, company, blueprint, orderlist = fetch_data()

date = '2021-06-01'
company_name = 'A'
blueprint_id = '1'
quantity = [(645001014, 1), (645001015, 2), (645001016, 3)]

company_info = search_company_info(company_name)
blueprint_info = search_blueprint_info(blueprint_id)
print("\n=== company_info ===\n")
print(date)
print(company[0][1])
print(company_info[0][1])
print(company[0][3])
print(blueprint_info[0][1])
print(company[0][4])
print("\n=== blueprint_info ===\n")


def add_with_tab(paragraph, left_text, left_value, right_text, right_value, tab_inches):
    tab_stop = paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(tab_inches))
    run = paragraph.add_run(left_text)
    run = paragraph.add_run(left_value)
    run.add_tab()
    paragraph.add_run(right_text)
    paragraph.add_run(right_value)
    
doc = Document()

head = doc.add_heading('발주서', level=0)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph0 = doc.add_paragraph()
add_with_tab(paragraph0, '견적번호: ', '1', '', '', 3)
paragraph1 = doc.add_paragraph()
add_with_tab(paragraph1, '발주일자: ', date, '본사: ', company[0][1], 3)
paragraph2 = doc.add_paragraph()
add_with_tab(paragraph2, '상호: ', company_info[0][1], '대표전화: ', company[0][3], 3)
paragraph3 = doc.add_paragraph()
add_with_tab(paragraph3, '공사명: ', blueprint_info[0][1], '팩스: ', company[0][4], 3)

table = doc.add_table(rows=1, cols=6) 
table.style = doc.styles['Table Grid']

#첫 행에 부품 정보 입력
first_row = table.rows[0].cells
first_row[0].text = '자재 번호'
first_row[1].text = '도면 번호'
first_row[2].text = '자재명'
first_row[3].text = '비고'
first_row[4].text = '수량'
first_row[5].text = '단가'

#문서에 요소 추가
def add_row(table, material):
    row_cells = table.add_row().cells
    row_cells[0].text = str(material[0][0])
    # row_cells[1].text = str(material['blueprint_number'])
    # row_cells[2].text = material['material_name']
    # row_cells[3].text = material['note']
    # row_cells[4].text = str(material['quantity'])
    # row_cells[5].text = str(material['price'])
row_cells = add_row(table, material)

doc.save('/Users/projects/database_web/database_web/test10.docx')